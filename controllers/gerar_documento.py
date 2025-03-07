from docx import Document
from docx.shared import Pt, Inches
import matplotlib.pyplot as plt
from io import BytesIO
from datetime import datetime
import os

def substituir_texto_em_elemento(elemento, placeholder, valor, fonte="Open Sans", tamanho=Pt(8)):
    valor = str(valor)
    texto_completo = elemento.text
    if placeholder in texto_completo:
        novo_texto = texto_completo.replace(placeholder, valor)
        elemento.clear()
        run = elemento.add_run(novo_texto)
        run.font.name = fonte
        run.font.size = tamanho
        return True
    else:
        for run in elemento.runs:
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, valor)
                run.font.name = fonte
                run.font.size = tamanho
                return True
    return False

def substituir_texto(doc, placeholder, valor, fonte="Open Sans", tamanho=Pt(8)):
    substituido = False
    for paragraph in doc.paragraphs:
        if substituir_texto_em_elemento(paragraph, placeholder, valor, fonte, tamanho):
            substituido = True
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if substituir_texto_em_elemento(paragraph, placeholder, valor, fonte, tamanho):
                        substituido = True
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            if substituir_texto_em_elemento(paragraph, placeholder, valor, fonte, tamanho):
                substituido = True
    return substituido

def gerar_tabela_como_imagem(tabelas_pesadas):
    fig, ax = plt.subplots(figsize=(6, 2))
    ax.axis("tight")
    ax.axis("off")
    colunas = ["Nome da Tabela", "Quantidade de Linhas", "Espaço Total (GB)"]
    dados = [[nome, "{:,}".format(linhas).replace(",", "."), espaco] for nome, (linhas, espaco) in tabelas_pesadas.items()]
    tabela = ax.table(cellText=dados, colLabels=colunas, loc="center", cellLoc="center")
    tabela.auto_set_font_size(False)
    tabela.set_fontsize(8)
    tabela.scale(1.2, 1.2)
    for (row, col), cell in tabela.get_celld().items():
        if row == 0:
            cell.set_facecolor("#D3D3D3")
            cell.set_text_props(weight="bold")
    buffer = BytesIO()
    plt.savefig(buffer, format="png", dpi=300, bbox_inches="tight")
    buffer.seek(0)
    plt.close()
    return buffer

def atualizar_tabela_como_imagem(doc, placeholder, tabelas_pesadas):
    for i, paragraph in enumerate(doc.paragraphs):
        if placeholder in paragraph.text:
            paragraph.clear()
            buffer = gerar_tabela_como_imagem(tabelas_pesadas)
            if buffer:
                run = paragraph.add_run()
                run.add_picture(buffer, width=Inches(5.0))
                buffer.close()
                return True
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if placeholder in paragraph.text:
                        paragraph.clear()
                        buffer = gerar_tabela_como_imagem(tabelas_pesadas)
                        if buffer:
                            run = paragraph.add_run()
                            run.add_picture(buffer, width=Inches(5.0))
                            buffer.close()
                            return True
    return False

def remover_tabela_banco(doc):
    tabela_encontrada = False
    inicio = None
    fim = None
    for i, table in enumerate(doc.tables):
        for row in table.rows:
            for cell in row.cells:
                if "{DATABASE_OU_USUARIO}" in cell.text:
                    tabela_encontrada = True
                    inicio = i
                if "{TABELAS_PESADAS}" in cell.text:
                    fim = i
                    break
    if tabela_encontrada and inicio is not None:
        if fim is None:
            fim = inicio
        for _ in range(fim, inicio - 1, -1):
            doc.element.body.remove(doc.tables[_]._element)

def gerar_documento(servidor, banco=None, tipo=None, dados_questionario=None):
    doc = Document("assets/Avaliação de Ambientes MODELO.docx")
    data_atual = datetime.now().strftime("%d/%m/%Y")
    substituir_texto(doc, "{DIA}", data_atual)

    if dados_questionario:
        substituir_texto(doc, "{ORDEM_SERVICO}", dados_questionario.get("ordem_servico", ""))
        substituir_texto(doc, "{NOME_TECNICO}", dados_questionario.get("nome_tecnico", ""))
        substituir_texto(doc, "{NOME_EMPRESA}", dados_questionario.get("nome_empresa", ""))
        substituir_texto(doc, "{NUMERO_USUARIOS}", dados_questionario.get("numero_usuarios", ""))
        substituir_texto(doc, "{VERSAO_CIGAM}", dados_questionario.get("versao_cigam", ""))
        substituir_texto(doc, "{HOSPEDAGEM}", dados_questionario.get("hospedagem", ""))

    if servidor:
        substituir_texto(doc, "{HOSTNAME}", servidor.hostname)
        substituir_texto(doc, "{SISTEMA_OPERACIONAL}", servidor.sistema_operacional)
        substituir_texto(doc, "{CPU}", servidor.cpu)
        substituir_texto(doc, "{ANO_CPU}", servidor.ano)
        substituir_texto(doc, "{NUCLEOS_CPU}", servidor.nucleos)
        substituir_texto(doc, "{THREADS_CPU}", servidor.threads)
        substituir_texto(doc, "{RAM}", servidor.ram)
        substituir_texto(doc, "{REDE}", servidor.rede)
        discos_text = "\n".join([f"{nome} - {info}" for nome, info in servidor.discos.items()])
        substituir_texto(doc, "{DISCOS}", discos_text)

    if banco:
        if tipo == "SQLServer":
            substituir_texto(doc, "{DATABASE_OU_USUARIO}", f"Database: {banco.nome_database}")
            substituir_texto(doc, "{VERSAO_BD}", banco.versao)
            substituir_texto(doc, "{TAMANHO_BD}", banco.datafile)
            substituir_texto(doc, "{TAMANHO_LOG}", "Tamanho do log")
            substituir_texto(doc, "{VALOR_TAMANHO_LOG}", banco.logfile)
            substituir_texto(doc, "{MEMORIA_MIN_OU_SGA}", "Memória mínima dedicada ao banco")
            substituir_texto(doc, "{VALOR_MEMORIA_MIN_OU_SGA}", banco.memoria_min)
            substituir_texto(doc, "{MEMORIA_MAX_OU_PGA}", "Memória máxima dedicada ao banco")
            substituir_texto(doc, "{VALOR_MEMORIA_MAX_OU_PGA}", banco.memoria_max)
        else:  # Oracle
            substituir_texto(doc, "{DATABASE_OU_USUARIO}", f"Usuário: {banco.usuario}")
            substituir_texto(doc, "{VERSAO_BD}", banco.versao)
            substituir_texto(doc, "{TAMANHO_BD}", banco.armazenamento)
            substituir_texto(doc, "{TAMANHO_LOG}", "")
            substituir_texto(doc, "{VALOR_TAMANHO_LOG}", "")
            substituir_texto(doc, "{MEMORIA_MIN_OU_SGA}", "SGA:")
            substituir_texto(doc, "{VALOR_MEMORIA_MIN_OU_SGA}", banco.sga)
            substituir_texto(doc, "{MEMORIA_MAX_OU_PGA}", "PGA:")
            substituir_texto(doc, "{VALOR_MEMORIA_MAX_OU_PGA}", banco.pga)
        atualizar_tabela_como_imagem(doc, "{TABELAS_PESADAS}", banco.tabelas_pesadas)
    else:
        remover_tabela_banco(doc)

    nome_empresa = dados_questionario.get("nome_empresa", "") if dados_questionario else ""
    nome_arquivo = f"Avaliação de Ambiente - {nome_empresa}.docx"
    doc.save(nome_arquivo)
    os.startfile(nome_arquivo)